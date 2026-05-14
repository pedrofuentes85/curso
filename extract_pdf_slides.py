#!/usr/bin/env python3
"""
Extrae texto de archivos PDF tipo diapositivas y genera un documento
consolidado en formato Word (.docx) o texto plano (.txt).

Uso:
    python extract_pdf_slides.py <dir_o_archivo.pdf> [<dir_o_archivo2.pdf> ...]
    python extract_pdf_slides.py diapositivas/ -o resumen -f docx
    python extract_pdf_slides.py *.pdf -o resumen -f txt
"""

import argparse
import os
import sys
from pathlib import Path


def extract_text_from_pdf(pdf_path: Path) -> list[tuple[int, str]]:
    """Extrae el texto de cada página de un PDF usando PyMuPDF.

    Returns:
        Lista de tuplas (numero_pagina, texto) para las páginas con contenido.
    """
    import fitz  # PyMuPDF

    pages = []
    with fitz.open(str(pdf_path)) as doc:
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if text.strip():
                pages.append((page_num, text.strip()))
    return pages


def generate_docx(
    pdf_data: list[tuple[str, list[tuple[int, str]]]], output_path: Path
) -> None:
    """Genera un documento Word con el texto extraído de los PDFs."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Estilo compacto: fuente pequeña para maximizar la densidad de información
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    for pdf_name, pages in pdf_data:
        # Encabezado por archivo
        heading = doc.add_heading(level=1)
        heading.clear()
        run = heading.add_run(pdf_name)
        run.font.size = Pt(12)

        for page_num, text in pages:
            # Etiqueta de diapositiva
            label = doc.add_paragraph()
            label_run = label.add_run(f"— Diapositiva {page_num} —")
            label_run.bold = True
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            # Contenido de la diapositiva
            doc.add_paragraph(text)

        doc.add_page_break()

    doc.save(str(output_path))


def generate_txt(
    pdf_data: list[tuple[str, list[tuple[int, str]]]], output_path: Path
) -> None:
    """Genera un archivo de texto plano con el texto extraído de los PDFs."""
    separator = "=" * 60

    with open(output_path, "w", encoding="utf-8") as f:
        for pdf_name, pages in pdf_data:
            f.write(f"\n{separator}\n")
            f.write(f"ARCHIVO: {pdf_name}\n")
            f.write(f"{separator}\n\n")

            for page_num, text in pages:
                f.write(f"--- Diapositiva {page_num} ---\n")
                f.write(text)
                f.write("\n\n")


def collect_pdf_files(inputs: list[str]) -> list[Path]:
    """Recopila todos los archivos PDF a partir de rutas de archivos o directorios."""
    pdf_files: list[Path] = []
    for input_path in inputs:
        p = Path(input_path)
        if p.is_file() and p.suffix.lower() == ".pdf":
            pdf_files.append(p)
        elif p.is_dir():
            pdf_files.extend(sorted(p.rglob("*.pdf")))
        else:
            print(
                f"Advertencia: '{input_path}' no es un PDF ni un directorio válido.",
                file=sys.stderr,
            )
    return pdf_files


def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Extrae texto de PDFs tipo diapositivas y genera un documento consolidado."
        )
    )
    parser.add_argument(
        "input",
        nargs="+",
        help="Archivos PDF o directorios que contienen PDFs",
    )
    parser.add_argument(
        "-o",
        "--output",
        default="contenido_extraido",
        help="Nombre del archivo de salida sin extensión (por defecto: contenido_extraido)",
    )
    parser.add_argument(
        "-f",
        "--format",
        choices=["docx", "txt"],
        default="docx",
        help="Formato de salida: 'docx' (Word, por defecto) o 'txt' (texto plano)",
    )

    args = parser.parse_args()

    # Recopilar archivos PDF
    pdf_files = collect_pdf_files(args.input)
    if not pdf_files:
        print("Error: No se encontraron archivos PDF.", file=sys.stderr)
        sys.exit(1)

    print(f"Procesando {len(pdf_files)} archivo(s) PDF...")

    # Extraer texto de cada PDF
    all_pdf_data: list[tuple[str, list[tuple[int, str]]]] = []
    for pdf_path in pdf_files:
        print(f"  Extrayendo: {pdf_path.name}")
        try:
            pages = extract_text_from_pdf(pdf_path)
            if pages:
                all_pdf_data.append((pdf_path.name, pages))
            else:
                print(
                    f"  Advertencia: No se encontró texto extraíble en '{pdf_path.name}'.",
                    file=sys.stderr,
                )
        except Exception as exc:
            print(
                f"  Error procesando '{pdf_path.name}': {exc}",
                file=sys.stderr,
            )

    if not all_pdf_data:
        print(
            "Error: No se pudo extraer texto de ningún PDF.",
            file=sys.stderr,
        )
        sys.exit(1)

    # Generar el archivo de salida
    output_path = Path(f"{args.output}.{args.format}")

    if args.format == "docx":
        print(f"Generando documento Word: {output_path}")
        generate_docx(all_pdf_data, output_path)
    else:
        print(f"Generando archivo de texto: {output_path}")
        generate_txt(all_pdf_data, output_path)

    size_kb = output_path.stat().st_size / 1024
    print(f"¡Listo! Archivo generado: {output_path} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
